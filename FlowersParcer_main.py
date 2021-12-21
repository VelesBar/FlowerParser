import requests
import docx
from docx import Document
from bs4 import BeautifulSoup as bs4
import os
from docx.enum.dml import MSO_THEME_COLOR_INDEX

'''
Основные функции
Получаем html
Парсим его, получаем ссылки и подписи
скачиваем картинки 
записываем всё в docx
'''

#Получаем html
def get_html(url):
    req = requests.get(url)
    soup = bs4(req.text, "html.parser")
    return soup

#получаем url для фильтров
def get_filters(url):
    req = requests.get(url)
    soup = bs4(req.text, "html.parser")
    filter_bar = soup.find('div', class_='sidebar_mobil')
    all_filters = filter_bar.findAll('div', class_='dropdown')

    #список названий фильтров, нужен для построения папок
    name_filters = []

    for i in range(len(all_filters)):
        name = all_filters[i].find('div', class_='text-block-25').text
        if name != 'Повод':
            name_filters.append(name)
    #список списков. Будем хранить ссылки по фильтрам
    any_list = []
    urls = []
    for i in range(len(all_filters)-1):
        urls.append([])

    #достаём фильтры по Стоимости
    filters_price=[]
    filters_price_url = all_filters[0].findAll('div', class_='block-input')
    for i in range(len(filters_price_url)):
        attribs = filters_price_url[i].find('input', type='radio').attrs['value']
        filters_price.append(attribs)
        urls[0].append('https://florartstudio.ru/catalog/'+'?minprice'+ attribs)

    #достаём фильтры по Категорям
    filters_flcat_url = all_filters[1].findAll('div', class_='block-input')
    filters_flcat = []
    for i in range(len(filters_flcat_url)):
        attribs = filters_flcat_url[i].find('input', type='checkbox').attrs['value']
        filters_flcat.append(attribs)
        urls[1].append('https://florartstudio.ru/catalog/'+'?flcat='+ attribs)

    #достаём фильтры по цветовой гамме
    filters_color_url = all_filters[3].findAll('div', class_='block-input')
    filters_color = []
    for i in range(len(filters_color_url)):
        attribs = filters_color_url[i].find('input', type='checkbox').attrs['value']
        filters_color.append(attribs)
        urls[2].append('https://florartstudio.ru/catalog/' + '?color=' + attribs)

    #достаём фильтры по "букеты с"
    filters_variety_url = all_filters[4].findAll('div', class_='block-input')
    filters_variety = []
    for i in range(len(filters_variety_url)):
        attribs = filters_variety_url[i].find('input', type='checkbox').attrs['value']
        filters_variety.append(attribs)
        urls[3].append('https://florartstudio.ru/catalog/' + '?variety=' + attribs)
    #собираем все фильтры
        filters = []
        filters.append(filters_price)
        filters.append(filters_flcat)
        filters.append(filters_color)
        filters.append(filters_variety)

    return urls, filters, name_filters


#Парсинг html. Получаем ссылки на картинки и подписи к ним
def html_parsing(html):
    page = html.find('div', class_="page-catalog__block")
    price_divs = page.findAll('div', class_='mini-card__price')
    page = page.findAll('a', class_='mini-card-info w-inline-block')

    images_link = [] #тут будут ссылки на карточки букетов
    images = [] #сюда запишу ссылки на картинки
    title = [] #тут буду хранить подписи к ним
    prices = [] #для цен
    links = [] #ссылки на букеты

    #Собираем
    for i in range(len(page)):
        images_link.append(page[i].attrs['href'])
        #подписи
        title.append(page[i].find('h3', class_='mini-card-header').text)
        #ссылки на картинки
        images.append(page[i].find('img', alt=title[i]).attrs['src'])
        #цены
        if price_divs[i].find('div', class_="mini-card-price-new__cena mini-card-price-new__cena_new") == None:
            prices.append(price_divs[i].find('div', class_="mini-card-price-new__cena").text)
        else:
            prices.append(price_divs[i].find('div', class_='mini-card-price-new__cena mini-card-price-new__cena_new').text)

    return images, title, prices, images_link

#функиця взята со StackOverflow
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

#скачиваем картинки. Необходимо вызвать до вызова create_docx!
def download_images(images_url, title):
    try:
        os.makedirs(r'images')
    except:
        print('Directory alrady exists')

    for i in range(len(images_url)):
        req = requests.get(images_url[i])
        with open(r'images\\' + title[i] + '.jpg', 'wb') as file:
            file.write(req.content)

def create_docx(images_url, titles, price, filter_name, dir_names):
    Docx = Document()
    for i in range(len(images_url)):
        Docx.add_paragraph(titles[i] + ' ' + price[i])
        p_link = Docx.add_paragraph()
        visible_link = images_and_imageslink[titles[i]]
        unvisible_link = images_and_imageslink[titles[i]]
        add_hyperlink(p_link,visible_link,unvisible_link)
        Docx.add_picture(r'images\\' + titles[i] + '.jpg')

    Docx.save(dir_names+'\\'+filter_name+'.docx')


def make_dir(dir_names):
    for i in range(len(dir_names)):
        try:
            os.makedirs(dir_names[i])
        except:
            print('Directory alrady exists')


"""
Краткий список функций и что они принимают/возвращают

def get_html(url):
def get_filters(url):
    return urls, filters, name_filters

def html_parsing(html):
    return images, title, prices, images_link

def download_images(images_url, title):
def create_docx(images_url, titles, price, filter_name):
"""
url = 'https://florartstudio.ru/catalog/'
filters = get_filters(url)
dir_names = filters[2]
urls = filters[0]
filters_title = filters[1]

images_and_titles = html_parsing(get_html(url))
images = images_and_titles[0]
titles = images_and_titles[1]
images_link = images_and_titles[3]
download_images(images, titles)
make_dir(dir_names)

#небольшой костыль, чтобы сопоставить ссылки на картинки с их именами. Нужно для корректной вставки в docx
images_and_imageslink = dict(zip(titles, images_link))



print(urls[0])
for i in range(len(urls)):
    for b in range(len(urls[i])):
        html = html_parsing(get_html(urls[i][b]))
        images = html[0]
        title = html[1]
        prices = html[2]
        create_docx(images, title, prices,filters_title[i][b],dir_names[i])
